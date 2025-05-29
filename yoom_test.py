import streamlit as st, pandas as pd
from io import BytesIO
from datetime import date, datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches
from PIL import Image
import re
import base64
import tempfile
import os
# from docx2pdf import convert # ModuleNotFoundError 해결을 위해 제거

# PDF 생성을 위한 추가 라이브러리 (reportlab 관련)
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics # make_pdf 함수 내부에 이미 import되어 있지만, 상단에 두어도 무방합니다.
from reportlab.pdfbase.ttfonts import TTFont # make_pdf 함수 내부에 이미 import되어 있지만, 상단에 두어도 무방합니다.

def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode()
    except:
        return None

# ────────────────────────────────────────────────
# 1) 페이지 설정
# ────────────────────────────────────────────────
st.set_page_config("경동나비엔 가스보일러 급배기전환 모델 확인 프로그램", layout="wide")

# 이미지 표시 방식 변경 (상단 중복 이미지 삭제)
# try:
#     st.image("images/kd.png", width=300)
# except:
#     st.error("이미지를 불러올 수 없습니다.")

# ────────────────────────────────────────────────
# 2) 세션 기본값
# ────────────────────────────────────────────────
def init_session_state():
    defaults = dict(
        page="model",
        status_html="",
        show_status=False,
        model_full="",
        qualification="",
        conversion_ok=False,
        판별완료=False,
        form_data={},
        history=[],
        # 첫 페이지 (model) 저장값
        selected_qualification="",
        # 두 번째 페이지 (product) 저장값
        selected_구분="",
        selected_세부구분="",
        selected_모델명="",
        selected_용량="",
        selected_연료="",
        selected_급배기방식="",
        # 세 번째 페이지 (form) 저장값
        form_번호="NO.1",
        form_연소기명="",
        form_수량=1,
        form_변경일자=date.today(),
        form_작업자_소속="",
        form_작업자_성명="",
        form_작업자격="가스보일러 제조사의 A/S 종사자",
        form_시공업체="",
        form_시공관리자=""
    )
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()
ss = st.session_state

# ────────────────────────────────────────────────
# 3) 보조 함수
# ────────────────────────────────────────────────
def capacity_ok(row, sel):
    if row["용량"].strip() == "없음":
        return sel == "없음"
    return sel in [c.strip() for c in row["용량"].split(",")]

def sanitize(name: str) -> str:          # ★ 파일명 안전 처리
    return re.sub(r'[\\/*?:"<>|]', "", name).strip() or "이름없음"


def make_docx(info: dict, sign_png: BytesIO | None) -> BytesIO:
    doc = Document()
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Pt(35))

    # 제목
    doc.add_paragraph("[별지 제44호 서식]<개정 23.07.11>").runs[0].font.size = Pt(10)
    p = doc.add_paragraph("연소기 변경 확인서")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    p = doc.add_paragraph("(제4-22조 및 제4-31조 관련)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # 기본 표
    tbl = doc.add_table(rows=3, cols=8)
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    h[0].text, h[1].text, h[2].text = "번호", "연소기명", "수량"
    h[3].text, h[4].text = "변경내역", "변경일자"
    h[5].merge(h[7]).text = "연소기 변경 작업자"
    sub = tbl.rows[1].cells
    sub[0].merge(tbl.cell(0, 0)); sub[1].merge(tbl.cell(0, 1)); sub[2].merge(tbl.cell(0, 2))
    sub[3].merge(tbl.cell(0, 3)); sub[4].merge(tbl.cell(0, 4))
    sub[5].text, sub[6].text, sub[7].text = "소 속", "성명(서명)", "작업자격"

    d = tbl.rows[2].cells
    d[0].text, d[1].text, d[2].text = info["번호"], info["연소기명"], str(info["수량"])
    d[3].text = "✔ 가스보일러 급배기방식 전환"
    d[4].text = info["변경일"].strftime("%Y-%m-%d")
    d[5].text, d[6].text, d[7].text = info["작업자_소속"], info["작업자_성명"], info["작업자격"]

    # 확인 문구
    doc.add_paragraph()
    doc.add_paragraph("상기와 같이 연소기 변경 작업을 실시하였음을 확인합니다.")
 # 날짜(우측 정렬)
    p_date = doc.add_paragraph(info["변경일"].strftime("%Y년 %m월 %d일"))
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 시공업체 줄 (우측 정렬)
    p_comp = doc.add_paragraph(f"○ 시공업체(상호): {info['시공업체']}")
    p_comp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 시공관리자 + 서명
    if sign_png:
        p_mgr = doc.add_paragraph()
        p_mgr.alignment = WD_ALIGN_PARAGRAPH.RIGHT      # ← 단락 정렬
        run = p_mgr.add_run(f"○ 시공관리자  : {info['시공관리자']}   (서명) ")
    else:
        p_mgr = doc.add_paragraph(f"○ 시공관리자  : {info['시공관리자']}   (서명) ")
        p_mgr.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # [비고] 표
    doc.add_paragraph()
    note_tbl = doc.add_table(rows=1, cols=1)
    note_tbl.style = "Table Grid"
    note = (
        "[비고]\n"
        "1. 변경내역은 해당되는 사항에 ✔ 표시\n"
        "2. 기술능력은 연소기 변경 작업자의 자격 기재\n"
        "   가. 열량법령 작업자격 : 지침 별표18 (예시 : 연소기 제조사 A/S 종사자)\n"
        "   나. 가스보일러 급배기방식 전환 작업자격 : KGS GC2008 또는 GC209 (예시 : 가스보일러 제조사 A/S 교육 이수자)"
    )
    note_tbl.cell(0, 0).text = note

    # docx 반환
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def make_pdf(info: dict) -> BytesIO:
    from io import BytesIO
    import os
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Table, TableStyle,
        Spacer, KeepTogether
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buffer = BytesIO()

    # 한글 폰트 등록
    # font_paths = [
    #     "C:/Windows/Fonts/malgun.ttf",
    #     "C:/Windows/Fonts/gulim.ttc",
    #     "C:/Windows/Fonts/batang.ttc",
    # ]
    # korean_font = None
    # for fp in font_paths:
    #     if os.path.exists(fp):
    #         try:
    #             pdfmetrics.registerFont(TTFont('Korean', fp))
    #             korean_font = 'Korean'
    #             break
    #         except:
    #             pass
    # if not korean_font:
    #     korean_font = 'Helvetica'

    # Streamlit Cloud 환경을 위해 시스템 폰트 사용 시도
    korean_font_name = 'UnDotum'
    try:
        pdfmetrics.registerFont(TTFont(korean_font_name, korean_font_name))
        korean_font = korean_font_name
    except:
        korean_font = 'Helvetica'

    # 문서 설정
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=20, leftMargin=20,
        topMargin=20, bottomMargin=20
    )

    # 스타일 생성 함수
    def make_style(name, size, align):
        return ParagraphStyle(
            name,
            fontName=korean_font,
            fontSize=size,
            leading=size * 1.6,
            alignment=align,
        )

    header_style   = make_style('Header',   9, TA_LEFT)
    title_style    = make_style('Title',   16, TA_CENTER)
    subtitle_style = make_style('Subt',    9, TA_CENTER)
    normal_style   = make_style('Normal',  9, TA_LEFT)
    right_style    = make_style('Right',   9, TA_RIGHT)

    story = []

    # 제목부
    story.append(Paragraph("[별지 제44호 서식] <개정 23.07.11>", header_style))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>연소기 변경 확인서</b>", title_style))
    story.append(Spacer(1, 4))
    story.append(Paragraph("(제4-22조 및 제4-31조 관련)", subtitle_style))
    story.append(Spacer(1, 12))

    # 표 데이터
    table_data = [
        ['번호','연소기명','수량','변경내역','변경일자','연소기 변경 작업자','',''],
        ['','','','','','소속','성명(서명)','작업자격'],
        [
            info['번호'],
            info['연소기명'],
            str(info['수량']),
            '✔ 가스보일러\n급배기방식\n전환',
            info['변경일'].strftime('%Y-%m-%d'),
            info['작업자_소속'],
            info['작업자_성명'],
            # 여기에 선택된 '작업자격'만 넣기
             info['작업자격'].replace(" ", "\n")  # 원하시면 공백을 줄바꿈으로 바꿀 수도 있습니다
    ]
]


    # 컬럼 폭 정의 (숫자로만)
    col_widths = [
        15* mm,  # 번호
        35* mm,  # 연소기명
        8 * mm,  # 수량
        30 * mm,  # 변경내역 (곱하기 연산자로 수정)
        25 * mm,  # 변경일자
        25 * mm,  # 소속
        25 * mm,  # 성명(서명)
        25 * mm,  # 작업자격
    ]

    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ('GRID',        (0,0), (-1,-1), 0.5, colors.black),
        ('FONTNAME',    (0,0), (-1,-1), korean_font),
        ('FONTSIZE',    (0,0), (-1,-1), 9),
        ('ALIGN',       (0,0), (-1,-1), 'CENTER'),
        ('VALIGN',      (0,0), (-1,-1), 'MIDDLE'),

        # "변경내역" 셀만 가로·세로 중앙정렬
        ('ALIGN',       (3,2), (3,2), 'CENTER'),
        ('VALIGN',      (3,2), (3,2), 'MIDDLE'),

        # 병합은 기존 그대로
        ('SPAN',        (0,0),(0,1)),
        ('SPAN',        (1,0),(1,1)),
        ('SPAN',        (2,0),(2,1)),
        ('SPAN',        (3,0),(3,1)),
        ('SPAN',        (4,0),(4,1)),
        ('SPAN',        (5,0),(7,0)),

        # 패딩 축소
        ('LEFTPADDING',  (0,0),(-1,-1), 2),
        ('RIGHTPADDING', (0,0),(-1,-1), 2),
        ('TOPPADDING',   (0,0),(-1,-1), 2),
        ('BOTTOMPADDING',(0,0),(-1,-1), 2),
    ]))
    
    # 확인 및 서명부
    confirm = Paragraph("상기와 같이 연소기 변경 작업을 실시하였음을 확인합니다.", normal_style)
    date_p  = Paragraph(info['변경일'].strftime('%Y년 %m월 %d일'), right_style)
    comp_p  = Paragraph(f"○ 시공업체(상호): {info['시공업체']}", right_style)
    mgr_p   = Paragraph(f"○ 시공관리자  : {info['시공관리자']}   (서명)", right_style)

        # ————————————————————————————————————————
    # 비고 표: HTML 태그로 줄바꿈·들여쓰기
    note_text = """
    <b>[비고]</b><br/>
    1. 변경내역은 해당되는 사항에 표시<br/>
    2. 기술능력은 연소기 변경 작업자의 자격 기재<br/>
    &nbsp;&nbsp;가. 열량법령 작업자격 : 지침 별표18 (예시 : 연소기 제조사 A/S 종사자)<br/>
    &nbsp;&nbsp;나. 가스보일러 급배기방식 전환 작업자격 : KGS GC2008 또는 GC209<br/>
    &nbsp;&nbsp;&nbsp;&nbsp;(예시 : 가스보일러 제조사 A/S 교육 이수자)
    """
    note_para = Paragraph(note_text, normal_style)

    note_table = Table([[note_para]], colWidths=[170*mm])
    note_table.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.5,colors.black),
        ('FONTNAME',(0,0),(-1,-1),korean_font),
        ('FONTSIZE',(0,0),(-1,-1),9),
        ('VALIGN',(0,0),(-1,-1),'TOP'),
        ('LEFTPADDING',(0,0),(-1,-1),4), ('RIGHTPADDING',(0,0),(-1,-1),4),
        ('TOPPADDING',(0,0),(-1,-1),4), ('BOTTOMPADDING',(0,0),(-1,-1),4),
    ]))

    # 한 페이지에 모두 묶기
    story.append(KeepTogether([
        table,
        Spacer(1,8),
        confirm,
        Spacer(1,8),
        date_p,
        Spacer(1,4),
        comp_p,
        Spacer(1,4),
        mgr_p,
        Spacer(1,12),
        note_table
    ]))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ────────────────────────────────────────────────
# 4) 데이터 (질문에서 주신 전체 리스트 그대로)
# ────────────────────────────────────────────────

data = [
    # ── 일반형 개방식 ──
    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB513", "연료": "LNG", "급배기방식": "FF",
     "용량": "13K, 16K, 20K, 25K, 30K, 35K", "비고": "대리점신축", "전환여부": "전환불가"},
    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB513", "연료": "LPG", "급배기방식": "FF",
     "용량": "13K, 16K, 20K, 25K, 30K, 35K", "비고": "대리점신축", "전환여부": "전환불가"},

    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB553", "연료": "LNG", "급배기방식": "FF",
     "용량": "13K, 16K, 20K, 25K, 30K, 35K", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB553", "연료": "LNG", "급배기방식": "FE",
     "용량": "13K, 16K, 20K, 25K, 30K, 35K", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB553", "연료": "LPG", "급배기방식": "FF",
     "용량": "13K, 16K", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB553", "연료": "LPG", "급배기방식": "FF",
     "용량": "20K, 25K, 30K, 35K", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "일반형", "세부구분": "개방식", "모델명": "NGB553", "연료": "LPG", "급배기방식": "FE",
     "용량": "20K, 25K, 30K, 35K", "비고": "대리점유통", "전환여부": "전환가능"},


    # ── 콘덴싱 개방식 ──
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB311", "연료": "LNG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K, 36K", "비고": "특판(단종예정)", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB311", "연료": "LPG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K, 36K", "비고": "특판(단종예정)", "전환여부": "전환불가"},
 
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB314", "연료": "LNG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "특판", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB314", "연료": "LPG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "특판", "전환여부": "전환불가"},

    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB324", "연료": "LNG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "대리점신축", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB324", "연료": "LPG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "대리점신축", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB354", "연료": "LNG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "대리점 유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB354", "연료": "LNG", "급배기방식": "FE",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "대리점 유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB354", "연료": "LPG", "급배기방식": "FF",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "대리점 유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB354", "연료": "LPG", "급배기방식": "FE",
     "용량": "15K, 18K, 22K, 27K, 33K", "비고": "대리점 유통", "전환여부": "전환가능"},


    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB384", "연료": "LNG", "급배기방식": "FF",
     "용량": "18K, 22K, 27K, 33K", "비고": "수요개발", "전환여부": "전환불가"},



    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB553", "연료": "LNG", "급배기방식": "FF",
     "용량": "22K, 27K, 33K, 43K", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB553", "연료": "LPG", "급배기방식": "FF",
     "용량": "22K, 27K, 33K, 43K", "비고": "대리점유통", "전환여부": "전환불가"},



    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB713", "연료": "LNG", "급배기방식": "FF",
     "용량": "22K, 27K, 33K, 43K", "비고": "특판", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB713", "연료": "LPG", "급배기방식": "FF",
     "용량": "22K, 27K, 33K, 43K", "비고": "특판", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB753", "연료": "LNG", "급배기방식": "FF",
     "용량": "22K, 27K, 33K, 43K", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "개방식", "모델명": "NCB753", "연료": "LPG", "급배기방식": "FF",
     "용량": "22K, 27K, 33K, 43K", "비고": "대리점유통", "전환여부": "전환불가"},



    # ── 일반형 밀폐식 ──

    {"구분": "일반형", "세부구분": "밀폐식", "모델명": "NGB553", "연료": "LNG", "급배기방식": "FF",
     "용량": "13L, 16L", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "일반형", "세부구분": "밀폐식", "모델명": "NGB553", "연료": "LNG", "급배기방식": "FF",
     "용량": "20L, 25L, 30L, 35L", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "일반형", "세부구분": "밀폐식", "모델명": "NGB553", "연료": "LNG", "급배기방식": "FE",
     "용량": "20L, 25L, 30L, 35L", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "일반형", "세부구분": "밀폐식", "모델명": "NGB553", "연료": "LPG", "급배기방식": "FF",
     "용량": "13L, 16L, 20L, 25L, 30L, 35L", "비고": "대리점유통", "전환여부": "전환불가"},


    # ── 콘덴싱 밀폐식 ──
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB311", "연료": "LNG", "급배기방식": "FF",
     "용량": "18L, 22L, 27L, 33L, 36L, 43L", "비고": "특판(단종예정)", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB311", "연료": "LPG", "급배기방식": "FF",
     "용량": "18L, 22L, 27L, 33L", "비고": "특판(단종예정)", "전환여부": "전환불가"},

    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB314", "연료": "LNG", "급배기방식": "FF",
     "용량": "18L, 22L, 27L, 33L", "비고": "특판", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB314", "연료": "LPG", "급배기방식": "FF",
     "용량": "18L, 22L, 27L, 33L", "비고": "특판", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB354", "연료": "LNG", "급배기방식": "FF",
     "용량": "15L, 18L, 22L, 27L, 33L", "비고": "대리점 유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB354", "연료": "LNG", "급배기방식": "FE",
     "용량": "15L, 18L, 22L, 27L, 33L", "비고": "대리점 유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB354", "연료": "LPG", "급배기방식": "FF",
     "용량": "15L, 18L, 22L, 27L, 33L", "비고": "대리점 유통", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB553", "연료": "LNG", "급배기방식": "FF",
     "용량": "22L, 27L, 33L", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB553", "연료": "LNG", "급배기방식": "FF",
     "용량": "43L", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB553", "연료": "LNG", "급배기방식": "FE",
     "용량": "43L", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB553", "연료": "LPG", "급배기방식": "FF",
     "용량": "22L, 27L, 33L, 43L", "비고": "대리점유통", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB713", "연료": "LNG", "급배기방식": "FF",
     "용량": "22L, 27L, 33L, 43L", "비고": "특판", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB713", "연료": "LPG", "급배기방식": "FF",
     "용량": "22L, 27L, 33L, 43L", "비고": "특판", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB753", "연료": "LNG", "급배기방식": "FF",
     "용량": "22L, 27L, 33L", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB753", "연료": "LNG", "급배기방식": "FF",
     "용량": "43L", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB753", "연료": "LNG", "급배기방식": "FE",
     "용량": "43L", "비고": "대리점유통", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB753", "연료": "LPG", "급배기방식": "FF",
     "용량": "22L, 27L, 33L, 43L", "비고": "대리점유통", "전환여부": "전환불가"},


    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB900", "연료": "LNG", "급배기방식": "FF",
     "용량": "43L, 52L", "비고": "대리점유통", "전환여부": "전환불가"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB900", "연료": "LPG", "급배기방식": "FF",
     "용량": "43L, 52L", "비고": "대리점유통", "전환여부": "전환불가"},

    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NPW(single)", "연료": "LNG", "급배기방식": "FF",
     "용량": "36KSS, 36KDS, 48KSS, 48KDS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NPW(single)", "연료": "LNG", "급배기방식": "FE",
     "용량": "36KSS, 36KDS, 48KSS, 48KDS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NPW(single)", "연료": "LPG", "급배기방식": "FF",
     "용량": "36KSS, 36KDS, 48KSS, 48KDS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NPW(single)", "연료": "LPG", "급배기방식": "FE",
     "용량": "36KSS, 36KDS, 48KSS, 48KDS", "비고": "단품용", "전환여부": "전환가능"},

    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LNG", "급배기방식": "FF",
     "용량": "45LSS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LNG", "급배기방식": "FE",
     "용량": "45LSS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LPG", "급배기방식": "FF",
     "용량": "45LSS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LPG", "급배기방식": "FE",
     "용량": "45LSS", "비고": "단품용", "전환여부": "전환가능"},

    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NFB790(single)", "연료": "LNG", "급배기방식": "FF",
     "용량": "75LSS, 100LSS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LNG", "급배기방식": "FE",
     "용량": "75LSS, 100LSS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LPG", "급배기방식": "FF",
     "용량": "75LSS, 100LSS", "비고": "단품용", "전환여부": "전환가능"},
    {"구분": "콘덴싱", "세부구분": "밀폐식", "모델명": "NCB790(single)", "연료": "LPG", "급배기방식": "FE",
     "용량": "75LSS, 100LSS", "비고": "단품용", "전환여부": "전환가능"},

    # ── 캐스케이드용 밀폐식 ──
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NPW", "연료": "LNG", "급배기방식": "FF",
     "용량": "36KS, 36KD, 48KS, 48KD", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NPW", "연료": "LNG", "급배기방식": "FE",
     "용량": "36KS, 36KD, 48KS, 48KD", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NPW", "연료": "LPG", "급배기방식": "FF",
     "용량": "36KS, 36KD, 48KS, 48KD", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NPW", "연료": "LPG", "급배기방식": "FE",
     "용량": "36KS, 36KD, 48KS, 48KD", "비고": "캐스케이드용", "전환여부": "전환가능"},

    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NCB790", "연료": "LNG", "급배기방식": "FF",
     "용량": "45LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NCB790", "연료": "LNG", "급배기방식": "FE",
     "용량": "45LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NCB790", "연료": "LPG", "급배기방식": "FF",
     "용량": "45LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NCB790", "연료": "LPG", "급배기방식": "FE",
     "용량": "45LS", "비고": "캐스케이드용", "전환여부": "전환가능"},

    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NFB790", "연료": "LNG", "급배기방식": "FF",
     "용량": "100LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NFB790", "연료": "LNG", "급배기방식": "FE",
     "용량": "100LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NFB790", "연료": "LPG", "급배기방식": "FF",
     "용량": "100LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
    {"구분": "캐스케이드용", "세부구분": "밀폐식", "모델명": "NFB790", "연료": "LPG", "급배기방식": "FE",
     "용량": "100LS", "비고": "캐스케이드용", "전환여부": "전환가능"},
]

df = pd.DataFrame(data)

# ────────────────────────────────────────────────
# 5) 페이지 로직
# ────────────────────────────────────────────────
if st.session_state.page == "model":
    st.title("경동나비엔 가스보일러 급배기전환 모델 확인 프로그램")

    # 설명과 이미지를 나란히 배치하기 위해 컬럼 사용
    col1, col2 = st.columns([1, 1])  # 왼쪽: 설명, 오른쪽: 이미지

    with col1:
        st.markdown("### 1. 급배기방식 전환 절차")
        st.markdown("""
        1) 가스보일러 포장박스 및 제품 명판에 **"본 제품은 급배기방식 (FF/FE) 전환이 가능합니다"** 라는 문구가 있는지 확인해주세요.  
        <span style='color:red;'>※ 제품 명판에 전환가능문구가 없으면 급배기전환 불가</span>  
        2) 전환설치가 가능한 제품은 설치설명서에 따라 작업해주세요.  
        3) 급배기방식 전환 후, 본체에 '급배기전환 표지판'을 부착해주세요.  
        4) '연소기 변경 확인서'를 검사처에 제출해주세요.  
        &nbsp;&nbsp;&nbsp;&nbsp;- 특정가스사용시설 또는 LPG 특정사용시설 등 → **안전공사검사원**에게 제출  
        &nbsp;&nbsp;&nbsp;&nbsp;- 특정가스사용시설 외 가정용보일러 설치시설 등 → **도시가스사**에 제출  
        """, unsafe_allow_html=True)

    with col2:
        # 이미지 표시
        try:
            st.image("images/kd.png", width=300)
        except:
            st.error("이미지를 불러올 수 없습니다.")

    # ✅ 여기서 col2 블록 벗어나 아래에 삽입
    st.markdown(
        "<h3>2. 급배기방식 전환 작업자의 자격 : 아래 항목 중 하나를 선택해주세요.</h3>",
        unsafe_allow_html=True
    )

    # 초기 선택값 설정
    if 'qualification_radio' not in st.session_state:
        st.session_state.qualification_radio = ss.selected_qualification if ss.selected_qualification else "가스보일러 제조사의 A/S 종사자"

    q = st.radio(
        "급배기전환 작업이 가능한 작업자인지 확인해주세요.",
        [
            "가스보일러 제조사의 A/S 종사자",
            "가스보일러 판매업체 직원으로서 가스보일러 제조사의 A/S 교육을 받은 자",
            "가스보일러 판매업체 직원으로서 A/S 업무에 2년 이상 근무한 자",
            "해당없음",
        ],
        key="qualification_radio",
        index=[
            "가스보일러 제조사의 A/S 종사자",
            "가스보일러 판매업체 직원으로서 가스보일러 제조사의 A/S 교육을 받은 자",
            "가스보일러 판매업체 직원으로서 A/S 업무에 2년 이상 근무한 자",
            "해당없음"
        ].index(st.session_state.qualification_radio)
    )
    ss.selected_qualification = st.session_state.qualification_radio

    # 자격 선택에 따른 메시지 표시 및 conversion_ok 상태 업데이트
    if q == "해당없음":
        st.markdown(
            '<p style="color:red;font-weight:bold;">※ 위 자격이 없는 설치업자는 급배기방식을 전환하여 설치할 수 없습니다.</p>',
            unsafe_allow_html=True,
        )
        ss.conversion_ok = False
    else:
        st.markdown(
            '<p style="color:blue;font-weight:bold;">◎ 급배기전환 작업이 가능합니다.</p>',
            unsafe_allow_html=True,
        )
        ss.conversion_ok = True # 자격 있으면 전환 가능으로 설정

    # '다음' 버튼 추가 (conversion_ok가 True일 때만 활성화)
    if st.button("다음", disabled=not ss.conversion_ok):
        ss.page = "product"
        # ss.판별완료 = True # 다음 페이지 이동 시 판별 완료 상태로 설정 (선택 사항, 필요시 주석 해제)
        st.rerun() # 페이지 전환을 위해 rerun 호출

# ────────────────────────────────────────────────
elif ss.page == "product":
    st.title("경동나비엔 가스보일러 급배기전환 모델 확인 프로그램")

    # 헤더 + '이전으로' 버튼 같이 표시
    head1, head2 = st.columns([4, 1])
    head1.markdown("### 급배기전환 제품을 선택하세요")
    if head2.button("◀ 이전으로", key="back_to_model"):
        ss.page = "model"
        ss.show_status = False     # 전환결과 숨기기
        ss['판별완료'] = False     # 확인서 버튼 비활성화
        ss.conversion_ok = False   # conversion 상태 초기화
        st.rerun()



    # 드롭다운 테두리 CSS 삽입
    st.markdown("""
    <style>
div[data-testid="stSelectbox"] > div {
    border: 1px solid black !important;
    border-radius: 4px !important;
    padding: 2px !important;
}
</style>
""", unsafe_allow_html=True)


    sel_g = st.selectbox("1. 구분", df["구분"].unique(), 
                        index=0 if not ss.selected_구분 else list(df["구분"].unique()).index(ss.selected_구분))
    ss.selected_구분 = sel_g
    df2 = df[df["구분"] == sel_g]

    # 세부구분 선택 로직 수정
    sub_category_list = list(df2["세부구분"].unique())
    sub_category_index = 0 if not ss.selected_세부구분 or ss.selected_세부구분 not in sub_category_list else sub_category_list.index(ss.selected_세부구분)
    sel_s = st.selectbox("2. 세부구분", sub_category_list,
                        index=sub_category_index)
    ss.selected_세부구분 = sel_s
    df3 = df2[df2["세부구분"] == sel_s]

    # 모델명 선택 로직 (이미 수정됨)
    model_list = list(df3["모델명"].unique())
    model_index = 0 if not ss.selected_모델명 or ss.selected_모델명 not in model_list else model_list.index(ss.selected_모델명)
    sel_m = st.selectbox("3. 모델명", model_list,
                        index=model_index)
    ss.selected_모델명 = sel_m

    df4 = df3[df3["모델명"] == sel_m]
    caps = []
    for cs in df4["용량"].unique():
        caps.extend(["없음"] if cs.strip() == "없음" else [c.strip() for c in cs.split(",")])
    caps = sorted(set(caps))

    # 용량 선택 로직 수정
    capacity_list = caps
    capacity_index = 0 if not ss.selected_용량 or ss.selected_용량 not in capacity_list else capacity_list.index(ss.selected_용량)
    sel_c = st.selectbox("4. 용량", capacity_list,
                        index=capacity_index)
    ss.selected_용량 = sel_c

    df5 = df4[df4.apply(lambda r: capacity_ok(r, sel_c), axis=1)]

    # 사용연료 선택 로직 수정
    fuel_list = list(df5["연료"].unique())
    fuel_index = 0 if not ss.selected_연료 or ss.selected_연료 not in fuel_list else fuel_list.index(ss.selected_연료)
    sel_f = st.selectbox("5. 사용연료", fuel_list,
                        index=fuel_index)
    ss.selected_연료 = sel_f
    df6 = df5[df5["연료"] == sel_f]

    # 급배기방식 선택 로직 수정
    exhaust_list = list(df6["급배기방식"].unique())
    exhaust_index = 0 if not ss.selected_급배기방식 or ss.selected_급배기방식 not in exhaust_list else exhaust_list.index(ss.selected_급배기방식)
    sel_v = st.selectbox("6. 급배기방식", exhaust_list,
                        index=exhaust_index)
    ss.selected_급배기방식 = sel_v


    # ── 판별 버튼 & 상태 메시지 + 버튼 같이 표시 ──
    btn_col, msg_col, form_col = st.columns([1, 3, 2])

    if '판별완료' not in ss:
        ss['판별완료'] = False

    if btn_col.button("판별하기"):
        fdf = df6[df6["급배기방식"] == sel_v]
        if fdf.empty:
            ss.show_status = False
            ss.conversion_ok = False
            ss['판별완료'] = False
            st.warning("선택한 조건에 맞는 모델이 없습니다. (또는 전환불가)")
        else:
            r = fdf.iloc[0]
            is_ok = "전환가능" in r["전환여부"]
            ss.conversion_ok = is_ok
            ss['판별완료'] = True

            status_text = "전환가능" if is_ok else "전환불가"
            word_html = (
                f'<span style="color:blue;font-weight:bold;">{status_text}</span>'
                if is_ok else
                f'<span style="color:red;font-weight:bold;">{status_text}</span>'
            )
            ss.status_html = word_html
            ss.show_status = True
            ss.model_full = f"{r['모델명']}-{sel_c} ({sel_f}, {sel_v})"

            sentence = (
                f"{r['비고']}에 설치되는 {r['구분']} 가스보일러 "
                f"{ss.model_full} ({r['세부구분']}) 는 급배기방식 {word_html} 합니다."
            )
            msg_col.markdown(sentence, unsafe_allow_html=True)

if ss.show_status:
    btn_col, msg_col, form_col = st.columns([1, 3, 2])  # 다시 선언
    if ss.conversion_ok:
        msg_col.markdown(
            f"""**전환여부 : {ss.status_html}**  
            <span style='font-size:0.9rem;'>(우측의 "연소기 변경 확인서 (급배기방식 전환)" 버튼을 눌러주세요)</span>
            """,
            unsafe_allow_html=True
        )
    else:
        msg_col.markdown(f"**전환여부 : {ss.status_html}**", unsafe_allow_html=True)

    form_col.button(
        "연소기 변경 확인서 (급배기방식 전환)",
        disabled=not (ss.get('판별완료') and ss.conversion_ok),
        on_click=lambda: (setattr(ss, "page", "form"), setattr(ss, "show_status", False))
    )

# ────────────────────────────────────────────────
elif ss.page == "form":
    st.title("연소기 변경 확인서 작성 (급배기방식 전환)")

    # ─── 상단에 '이전' 버튼 추가 ───
    if st.button("◀ 이전으로", key="back_to_product"):
        ss.page = "product"
        ss.show_status = True  # 전환결과 표시 유지
        st.rerun()

    # == 상단 : 제품 정보 ==
    st.markdown("### ■ 급배기전환 제품 정보")
    g1, g2, g3, g4 = st.columns([1, 3, 1, 1])
    번호 = g1.text_input("번호", value=ss.form_번호, disabled=True, label_visibility="collapsed")
    연소기명 = g2.text_input("연소기명", value=ss.form_연소기명 or ss.model_full, disabled=True, label_visibility="collapsed")
    수량 = g3.number_input("수량", min_value=1, value=ss.form_수량, label_visibility="collapsed")
    변경일자 = g4.date_input("변경일자", value=ss.form_변경일자, label_visibility="collapsed")

    # 입력값 저장
    ss.form_번호 = 번호
    ss.form_연소기명 = 연소기명
    ss.form_수량 = 수량
    ss.form_변경일자 = 변경일자

    # 라벨 표시를 별도 줄에 배치
    g1.caption("번호"); g2.caption("연소기명"); g3.caption("수량"); g4.caption("변경일자")

    st.checkbox("가스보일러 급배기방식 전환 ", value=True, disabled=True)

    # == 작업자 정보 ==
    st.markdown("### ■ 연소기 변경 작업자 정보")
    j1, j2, j3 = st.columns([1, 1, 2])
    작업자_소속 = j1.text_input("소속", value=ss.form_작업자_소속)
    작업자_성명 = j2.text_input("성명(서명)", value=ss.form_작업자_성명)
    radio = [
        "가스보일러 제조사의 A/S 종사자",
        "가스보일러 판매업체 직원으로서 제조사 A/S 교육 이수자",
        "가스보일러 판매업체 직원으로서 A/S 업무 2년 이상",
    ]
    작업자격 = j3.radio("작업자격", radio, 
                    index=0 if not ss.form_작업자격 else radio.index(ss.form_작업자격))

    # 입력값 저장
    ss.form_작업자_소속 = 작업자_소속
    ss.form_작업자_성명 = 작업자_성명
    ss.form_작업자격 = 작업자격

    s1, s2 = st.columns(2)
    시공업체 = s1.text_input("시공업체(상호)", value=ss.form_시공업체)
    시공관리자 = s2.text_input("시공관리자", value=ss.form_시공관리자)

    # 입력값 저장
    ss.form_시공업체 = 시공업체
    ss.form_시공관리자 = 시공관리자

    # ★ 바로 아래 이 위치에 CSS 추가하세요!
    st.markdown("""
    <style>
    /* 수량, 변경일자: 전체 컨테이너 div에 테두리 적용 */
    div[data-testid="stNumberInput"] {
        border: 2px solid black !important;
        border-radius: 6px !important;
        padding: 4px;
    }
    div[data-testid="stDateInput"] {
        border: 2px solid black !important;
        border-radius: 6px !important;
        padding: 4px;
    }

    /* 텍스트 입력창은 여전히 input 요소에 적용 */
    div[data-testid="stTextInput"] input[aria-label="소속"],
    div[data-testid="stTextInput"] input[aria-label="성명(서명)"],
    div[data-testid="stTextInput"] input[aria-label="시공업체(상호)"],
    div[data-testid="stTextInput"] input[aria-label="시공관리자"] {
        border: 2px solid black !important;
        border-radius: 6px !important;
        padding: 4px;
    }
    </style>
    """, unsafe_allow_html=True)
    # ── 다운로드 버튼 ──
    if st.button("연소기 변경 확인서 다운로드"):
        try:
            # 필수 입력값 검증
            if not all([작업자_소속, 작업자_성명, 시공업체, 시공관리자]):
                st.error("모든 필수 항목을 입력해주세요.")
                st.stop()

            # 현재 입력 정보를 딕셔너리로 저장
            current_data = {
                "번호": "NO.1",
                "연소기명": 연소기명,
                "수량": 수량,
                "변경일": 변경일자,
                "작업자_소속": 작업자_소속,
                "작업자_성명": 작업자_성명,
                "작업자격": 작업자격,
                "시공업체": 시공업체,
                "시공관리자": 시공관리자,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            # 히스토리에 추가
            st.session_state.history.append(current_data)

            # 파일명 기본 부분
            base_name = f"연소기_변경_확인서_{sanitize(시공관리자)}"
            
            # 문서 정보
            doc_info = dict(
                번호="NO.1", 
                연소기명=연소기명, 
                수량=수량, 
                변경일=변경일자,
                작업자_소속=작업자_소속, 
                작업자_성명=작업자_성명, 
                작업자격=작업자격,
                시공업체=시공업체, 
                시공관리자=시공관리자
            )

            # 두 개의 버튼을 나란히 배치
            col1, col2 = st.columns(2)

            with col1:
                # 워드 파일 다운로드
                word_buf = make_docx(doc_info, None)
                st.download_button(
                    "📄 Word 파일 저장",
                    data=word_buf.getvalue(),
                    file_name=f"{base_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word"
                )

            with col2:
                # PDF 파일 다운로드 (make_pdf 함수 사용)
                try:
                    pdf_buf = make_pdf(doc_info)
                    st.download_button(
                        "📄 PDF 파일 저장",
                        data=pdf_buf.getvalue(),
                        file_name=f"{base_name}.pdf",
                        mime="application/pdf",
                        key="download_pdf"
                    )
                except Exception as e:
                    st.error(f"PDF 생성 중 오류가 발생했습니다: {str(e)}")
                    st.error("ReportLab 관련 오류일 수 있습니다. 필요한 라이브러리가 설치되었는지 확인해주세요.")

        except Exception as e:
            # 전체 문서 생성 오류 처리
            st.error(f"문서 생성 중 오류가 발생했습니다: {str(e)}")
            st.error("필수 입력 항목을 다시 확인하거나 잠시 후 다시 시도해주세요.")